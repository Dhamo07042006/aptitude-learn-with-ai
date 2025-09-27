import os
import pandas as pd
import matplotlib
matplotlib.use("Agg")  # Non-GUI backend for server-side plotting
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches
import random

# ------------------ Enhanced Rule-based AI Summary ------------------
def generate_ai_summary(df):
    """
    Generate a professional, polished summary paragraph with performance analysis and motivation.
    No student name repeated in the paragraph.
    """
    total_questions = len(df)
    correct_count = df['is_correct'].sum()
    accuracy = (correct_count / total_questions) * 100 if total_questions > 0 else 0
    avg_time = df['time_taken'].mean() if 'time_taken' in df else 0

    # Identify strengths and weaknesses
    topic_acc = df.groupby("topic")["is_correct"].mean() * 100
    strengths = [topic for topic, acc in topic_acc.items() if acc >= 75]
    weaknesses = [topic for topic, acc in topic_acc.items() if acc < 75]

    subtopic_acc = df.groupby("subtopic")["is_correct"].mean() * 100
    weak_subtopics = [sub for sub, acc in subtopic_acc.items() if acc < 60]

    # Professional phrasing templates
    intro_templates = [
        "The performance in this assessment demonstrates a diligent approach to learning.",
        "This report highlights the results across various topics and subtopics, reflecting strengths and areas for improvement.",
        "Engagement and performance metrics provide insight into overall academic progress."
    ]
    strength_templates = [
        f"Strong performance was observed in: {', '.join(strengths)}." if strengths else "",
        f"Consistent accuracy was achieved in topics such as: {', '.join(strengths)}." if strengths else ""
    ]
    improvement_templates = [
        f"Topics requiring further attention include: {', '.join(weaknesses)}." if weaknesses else "",
        f"Areas that may benefit from additional practice: {', '.join(weak_subtopics)}." if weak_subtopics else ""
    ]
    motivation_templates = [
        "Focused practice on weaker areas is recommended to achieve higher accuracy in future assessments.",
        "Maintaining consistent effort and reviewing challenging topics will lead to better outcomes.",
        "Optimizing time management and concentrating on improvement areas can enhance overall performance."
    ]

    # Construct paragraph
    paragraph = " ".join(random.sample(intro_templates, 1))
    if any(strength_templates):
        paragraph += " " + " ".join(filter(None, strength_templates))
    if any(improvement_templates):
        paragraph += " " + " ".join(filter(None, improvement_templates))
    paragraph += f" Overall accuracy: {accuracy:.2f}%, Average time per question: {avg_time:.2f}s. "
    paragraph += random.choice(motivation_templates)
    return paragraph

# ------------------ Report Generation ------------------
def generate_report(solutions, output_dir="reports", student_name="Student"):
    """
    Generate a Word report with charts, question-wise analysis, 
    result analysis, and a professional AI-generated motivational paragraph.
    Student name mentioned only once.
    """
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    df = pd.DataFrame(solutions)
    df["is_correct"] = df["is_correct"].astype(bool)
    df["time_taken"] = df["time_taken"].astype(float)

    # ----- Topic Accuracy -----
    topic_acc = df.groupby("topic")["is_correct"].mean() * 100
    plt.figure(figsize=(8,5))
    sns.barplot(x=topic_acc.index, y=topic_acc.values, palette="Blues_d")
    plt.ylim(0,100)
    plt.ylabel("Accuracy (%)")
    plt.title("Topic Accuracy")
    topic_chart_path = os.path.join(output_dir, "topic_accuracy.png")
    plt.savefig(topic_chart_path)
    plt.close()

    # ----- Subtopic Accuracy -----
    subtopic_acc = df.groupby("subtopic")["is_correct"].mean() * 100
    plt.figure(figsize=(10,5))
    sns.barplot(x=subtopic_acc.index, y=subtopic_acc.values, palette="Greens_d")
    plt.ylim(0,100)
    plt.ylabel("Accuracy (%)")
    plt.title("Subtopic Accuracy")
    subtopic_chart_path = os.path.join(output_dir, "subtopic_accuracy.png")
    plt.savefig(subtopic_chart_path)
    plt.close()

    # ----- Learning Fundamentals -----
    fundamentals = {
        "Listening": max(0, 100 - df["time_taken"].mean()),
        "Grasping": df["is_correct"].mean() * 100,
        "Retention": 100 - df[~df["is_correct"]].shape[0]/df.shape[0]*100,
        "Application": df[df.get("difficulty", pd.Series(["Very easy"]*len(df))).isin(["Moderate","Difficult"])]["is_correct"].mean() * 100
    }

    plt.figure(figsize=(6,4))
    sns.barplot(x=list(fundamentals.keys()), y=list(fundamentals.values()), palette="Oranges_d")
    plt.ylim(0,100)
    plt.ylabel("Score (%)")
    plt.title("Learning Fundamentals")
    fundamentals_chart_path = os.path.join(output_dir, "fundamentals.png")
    plt.savefig(fundamentals_chart_path)
    plt.close()

    # ----- Generate Word Report -----
    report_path = os.path.join(output_dir, f"{student_name}_report.docx")
    doc = Document()
    doc.add_heading(f"{student_name} - Performance Report", 0)
    doc.add_paragraph("This report provides detailed insights into performance by topic, subtopic, and learning fundamentals.\n")

    # Learning Fundamentals
    doc.add_heading("Learning Fundamentals", level=1)
    for k, v in fundamentals.items():
        doc.add_paragraph(f"{k}: {v:.2f}%")
    doc.add_picture(fundamentals_chart_path, width=Inches(5))

    # Topic Accuracy
    doc.add_heading("Topic Accuracy", level=1)
    doc.add_picture(topic_chart_path, width=Inches(5))

    # Subtopic Accuracy
    doc.add_heading("Subtopic Accuracy", level=1)
    doc.add_picture(subtopic_chart_path, width=Inches(5))

    # ----- Result Analysis -----
    doc.add_heading("Result Analysis", level=1)
    total_questions = len(df)
    correct_count = df['is_correct'].sum()
    incorrect_count = total_questions - correct_count
    avg_time = df['time_taken'].mean()
    doc.add_paragraph(f"Total Questions: {total_questions}")
    doc.add_paragraph(f"Correct Answers: {correct_count}")
    doc.add_paragraph(f"Incorrect Answers: {incorrect_count}")
    doc.add_paragraph(f"Average Time per Question: {avg_time:.2f} seconds")

    # Question-wise Performance
    doc.add_heading("Question-wise Performance", level=1)
    for i, row in df.iterrows():
        doc.add_paragraph(
            f"Q: {row['question']}\n"
            f"Topic/Subtopic: {row['topic']} / {row['subtopic']}\n"
            f"Your Answer: {row['user_answer']} | Correct Answer: {row['correct_answer']} | "
            f"{'✅ Correct' if row['is_correct'] else '❌ Incorrect'} | Time Taken: {row['time_taken']}s"
        )

    # ----- AI-generated Summary & Motivation -----
    doc.add_heading("AI Analysis", level=1)
    ai_text = generate_ai_summary(df)
    doc.add_paragraph(ai_text)

    doc.save(report_path)
    print(f"Report generated: {report_path}")
    return report_path
